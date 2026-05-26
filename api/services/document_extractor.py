"""Document text extraction service for PDF and DOCX files.

Handles various document edge cases with appropriate error messages in Spanish.
"""

from dataclasses import dataclass
from io import BytesIO
from typing import BinaryIO

import fitz  # PyMuPDF
from docx import Document as DocxDocument


class DocumentExtractionError(Exception):
    """Custom exception for document extraction errors."""

    def __init__(self, message: str, code: str):
        self.message = message
        self.code = code
        super().__init__(message)


# Alias for backward compatibility
PDFExtractionError = DocumentExtractionError


@dataclass
class ExtractionResult:
    """Result of document text extraction."""

    text: str
    page_count: int
    char_count: int
    file_type: str  # "pdf" or "docx"


# Error messages in Spanish
ERRORS = {
    "oversized": "Archivo muy grande (máx 5MB)",
    "protected": "PDF protegido, removí la contraseña",
    "corrupt": "Archivo corrupto o inválido",
    "image_only": "PDF sin texto legible (escaneado?)",
    "too_long": "CV muy largo (máx 10 páginas)",
    "insufficient": "CV sin contenido suficiente",
    "unsupported": "Formato no soportado (solo PDF y DOCX)",
}

# Limits
MAX_FILE_SIZE = 5 * 1024 * 1024  # 5MB
MAX_PAGES = 10
MIN_CHARS_PER_PAGE = 100  # Minimum average chars per page to not be "image only"
MIN_TOTAL_CHARS = 200  # Minimum total chars for meaningful content


def extract_text(
    file_content: bytes | BinaryIO,
    filename: str = "cv.pdf",
) -> ExtractionResult:
    """Extract text content from a PDF or DOCX file.

    Args:
        file_content: File bytes or file-like object
        filename: Original filename (used to detect file type)

    Returns:
        ExtractionResult with extracted text and metadata

    Raises:
        DocumentExtractionError: If the document cannot be processed
    """
    # Convert to bytes if needed
    if hasattr(file_content, "read"):
        file_bytes = file_content.read()
    else:
        file_bytes = file_content

    # Check file size
    if len(file_bytes) > MAX_FILE_SIZE:
        raise DocumentExtractionError(ERRORS["oversized"], "oversized")

    # Detect file type by extension
    filename_lower = filename.lower()
    if filename_lower.endswith(".pdf"):
        return _extract_from_pdf(file_bytes)
    elif filename_lower.endswith(".docx"):
        return _extract_from_docx(file_bytes)
    else:
        raise DocumentExtractionError(ERRORS["unsupported"], "unsupported")


def _extract_from_pdf(file_bytes: bytes) -> ExtractionResult:
    """Extract text from PDF file."""
    try:
        doc = fitz.open(stream=file_bytes, filetype="pdf")
    except fitz.fitz.FileDataError:
        raise DocumentExtractionError(ERRORS["corrupt"], "corrupt")
    except Exception as e:
        if "password" in str(e).lower() or "encrypted" in str(e).lower():
            raise DocumentExtractionError(ERRORS["protected"], "protected")
        raise DocumentExtractionError(ERRORS["corrupt"], "corrupt")

    try:
        if doc.is_encrypted:
            raise DocumentExtractionError(ERRORS["protected"], "protected")

        page_count = len(doc)
        if page_count > MAX_PAGES:
            raise DocumentExtractionError(ERRORS["too_long"], "too_long")

        if page_count == 0:
            raise DocumentExtractionError(ERRORS["corrupt"], "corrupt")

        text_parts: list[str] = []
        for page in doc:
            page_text = page.get_text("text")
            if page_text:
                text_parts.append(page_text.strip())

        full_text = "\n\n".join(text_parts)
        char_count = len(full_text)

        avg_chars_per_page = char_count / page_count if page_count > 0 else 0
        if avg_chars_per_page < MIN_CHARS_PER_PAGE:
            raise DocumentExtractionError(ERRORS["image_only"], "image_only")

        if char_count < MIN_TOTAL_CHARS:
            raise DocumentExtractionError(ERRORS["insufficient"], "insufficient")

        return ExtractionResult(
            text=full_text,
            page_count=page_count,
            char_count=char_count,
            file_type="pdf",
        )

    finally:
        doc.close()


def _extract_from_docx(file_bytes: bytes) -> ExtractionResult:
    """Extract text from DOCX file."""
    try:
        doc = DocxDocument(BytesIO(file_bytes))
    except Exception:
        raise DocumentExtractionError(ERRORS["corrupt"], "corrupt")

    try:
        # Extract text from paragraphs
        paragraphs: list[str] = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)

        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    paragraphs.append(row_text)

        full_text = "\n\n".join(paragraphs)
        char_count = len(full_text)

        # Estimate page count (roughly 2500 chars per page for Word docs)
        estimated_pages = max(1, char_count // 2500 + 1)

        if estimated_pages > MAX_PAGES:
            raise DocumentExtractionError(ERRORS["too_long"], "too_long")

        if char_count < MIN_TOTAL_CHARS:
            raise DocumentExtractionError(ERRORS["insufficient"], "insufficient")

        return ExtractionResult(
            text=full_text,
            page_count=estimated_pages,
            char_count=char_count,
            file_type="docx",
        )

    except DocumentExtractionError:
        raise
    except Exception:
        raise DocumentExtractionError(ERRORS["corrupt"], "corrupt")


# Backward compatibility aliases
def extract_text_from_pdf(
    file_content: bytes | BinaryIO,
    filename: str = "cv.pdf",
) -> ExtractionResult:
    """Backward compatible function for PDF extraction."""
    return extract_text(file_content, filename)


def validate_cv_content(text: str) -> None:
    """Validate that extracted text has enough content for analysis."""
    if len(text.strip()) < MIN_TOTAL_CHARS:
        raise DocumentExtractionError(ERRORS["insufficient"], "insufficient")
